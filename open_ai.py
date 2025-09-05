from flask import Flask, request, jsonify, send_file, render_template, redirect
from flask_cors import CORS
import boto3
import os
import random
import string
import io
import qrcode
import time
import threading
from dotenv import load_dotenv
from openai import OpenAI
from PyPDF2 import PdfReader, PdfMerger
from docx import Document
from reportlab.pdfgen import canvas
from PIL import Image
from pdf2image import convert_from_bytes
import pytesseract

# ----------------------- App Setup -----------------------
app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}})
app.config['MAX_CONTENT_LENGTH'] = 512 * 1024 * 1024  # 512 MB

UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
DELETE_AFTER = 2 * 60 * 60
CHECK_INTERVAL = 300

# ----------------------- Load Secrets -----------------------
load_dotenv()
aws_access_key = os.environ.get("AWS_Access_key")
aws_secret_key = os.environ.get("AWS_secret_key")
openai_api_key = os.environ.get("OPENAI_API_KEY")
BUCKET_NAME = os.environ.get("BUCKET_NAME", "ycloud45")  # allow override via env

# Optional tesseract path from env (if installed in custom location)
tess_cmd = os.environ.get("TESSERACT_CMD")
if tess_cmd:
    pytesseract.pytesseract.tesseract_cmd = tess_cmd

# Validate critical envs
if not openai_api_key:
    print("Warning: OPENAI_API_KEY not set. AI features will fail until provided.")

s3 = boto3.client(
    's3',
    aws_access_key_id=aws_access_key,
    aws_secret_access_key=aws_secret_key,
    region_name='us-east-1'
)

# ----------------------- OpenAI Client -----------------------
client = OpenAI(api_key=openai_api_key)


def upload_bytes_to_s3(file_bytes, key, content_type="application/octet-stream"):
    """
    Upload bytes to S3 and return a presigned GET URL (valid 1 hour) or None on failure.
    """
    try:
        s3.put_object(Bucket=BUCKET_NAME, Key=key, Body=file_bytes, ContentType=content_type)
        url = s3.generate_presigned_url(
            'get_object',
            Params={'Bucket': BUCKET_NAME, 'Key': key},
            ExpiresIn=3600
        )
        return url
    except Exception as e:
        print("S3 Upload Error:", e)
        return None


def generate_code(length=6):
    return ''.join(random.choices(string.ascii_uppercase + string.digits, k=length))


def auto_delete():
    while True:
        now = time.time()
        if os.path.exists(UPLOAD_FOLDER):
            for i in os.listdir(UPLOAD_FOLDER):
                path = os.path.join(UPLOAD_FOLDER, i)
                if os.path.isfile(path) and now - os.path.getmtime(path) > DELETE_AFTER:
                    try:
                        os.remove(path)
                    except Exception as e:
                        print("Error deleting file:", e)
        time.sleep(CHECK_INTERVAL)


threading.Thread(target=auto_delete, daemon=True).start()

# ----------------------- Routes -----------------------
@app.route("/")
def home():
    return render_template("index.html")


@app.route("/upload", methods=["POST"])
def upload():
    """
    Returns a presigned PUT URL for direct client upload.
    Expects JSON body with "filename" and "filetype" (content-type) keys.
    """
    data = request.get_json(force=True, silent=True)
    if not data or "filename" not in data or "filetype" not in data:
        return jsonify({"error": "filename and filetype are required in JSON body"}), 400

    code = generate_code()
    filename = f"{code}_{data['filename']}"

    presigned_url = None
    try:
        presigned_url = s3.generate_presigned_url(
            'put_object',
            Params={
                'Bucket': BUCKET_NAME,
                'Key': filename,
                'ContentType': data['filetype']
            },
            ExpiresIn=600
        )
    except Exception as e:
        print("Presigned URL generation error:", e)
        return jsonify({"error": "Could not generate upload URL"}), 500

    return jsonify({
        "code": code,
        "upload_url": presigned_url
    }), 200


@app.route("/api/status", methods=["GET"])
def api_status():
    return jsonify({
        "status": "running",
        "timestamp": time.time()
    }), 200


@app.route("/download/<code>", methods=["GET"])
def download(code):
    code = code.strip()
    print("code:", code)
    try:
        res = s3.list_objects_v2(Bucket=BUCKET_NAME, Prefix=f"{code}_")
    except Exception as e:
        print("S3 list error:", e)
        return jsonify({"error": "S3 error"}), 500

    contents = res.get("Contents") or []
    if not contents:
        print("No file found for code:", code)
        return jsonify({"error": "Invalid code or file not found"}), 404

    file_key = contents[0]["Key"]
    print("Matched_file:", file_key)
    try:
        url = s3.generate_presigned_url(
            "get_object",
            Params={"Bucket": BUCKET_NAME, "Key": file_key, "ResponseContentDisposition": "attachment"},
            ExpiresIn=3600
        )
    except Exception as e:
        print("Presigned GET error:", e)
        return jsonify({"error": "Could not generate download URL"}), 500

    return redirect(url, code=302)


@app.route('/get_qr/<code>', methods=['GET'])
def get_qr(code):
    download_url = f"{request.scheme}://{request.host}/download/{code}"
    qr = qrcode.make(download_url)
    buf = io.BytesIO()
    qr.save(buf, format="PNG")
    buf.seek(0)
    return send_file(buf, mimetype='image/png')


def upload_to_s3(file_content, file_name, folder=""):
    key = f"{folder}/{file_name}" if folder else file_name
    s3.put_object(Bucket=BUCKET_NAME, Key=key, Body=file_content)
    return f"s3://{BUCKET_NAME}/{key}"


# ----------------------- Document Helpers -----------------------
def extract(file_like):
    """
    Accepts a file-like object (supports .read()). Returns extracted text (PDF).
    """
    try:
        # if it's a Flask FileStorage, read bytes
        if hasattr(file_like, "read"):
            data = file_like.read()
        else:
            data = file_like
        reader = PdfReader(io.BytesIO(data))
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text
    except Exception as e:
        print(f"PDF extraction error: {e}")
        return ""


def ask_ai(prompt, max_tokens=500):
    """Send prompt to OpenAI and return response (best-effort)."""
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a helpful and cute assistant"},
                {"role": "user", "content": prompt}
            ],
            max_tokens=max_tokens
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        print("OpenAI error:", e)
        return ""


def pdf_to_word_bytes(pdf_bytes):
    try:
        reader = PdfReader(io.BytesIO(pdf_bytes))
        doc = Document()
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                doc.add_paragraph(f"--- Page {i + 1} ---")
                doc.add_paragraph(text)
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return f.getvalue()
    except Exception as e:
        print(f"PDF to Word conversion error: {e}")
        return None


def word_to_pdf_bytes(word_bytes):
    try:
        doc = Document(io.BytesIO(word_bytes))
        f = io.BytesIO()
        c = canvas.Canvas(f)
        y = 800
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                c.drawString(50, y, text)
                y -= 20
                if y < 50:
                    c.showPage()
                    y = 800
        c.save()
        f.seek(0)
        return f.getvalue()
    except Exception as e:
        print(f"Word to PDF conversion error: {e}")
        return None


def image_to_pdf_bytes(image_bytes):
    try:
        img = Image.open(io.BytesIO(image_bytes))
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")
        f = io.BytesIO()
        img.save(f, "PDF")
        f.seek(0)
        return f.getvalue()
    except Exception as e:
        print(f"Image to PDF conversion error: {e}")
        return None


def pdf_to_images_bytes(pdf_bytes):
    try:
        pages = convert_from_bytes(pdf_bytes, 300)
        images_bytes = []
        for i, page in enumerate(pages):
            buf = io.BytesIO()
            page.save(buf, format="PNG")
            buf.seek(0)
            images_bytes.append(buf.getvalue())
        return images_bytes
    except Exception as e:
        print(f"PDF to images conversion error: {e}")
        return []


def merge_pdf_bytes(pdf_files_bytes):
    try:
        merger = PdfMerger()
        for pdf_bytes in pdf_files_bytes:
            merger.append(io.BytesIO(pdf_bytes))
        out = io.BytesIO()
        merger.write(out)
        merger.close()
        out.seek(0)
        return out.getvalue()
    except Exception as e:
        print(f"PDF merge error: {e}")
        return None


def resize_image(image_bytes, max_width=800, max_height=800):
    try:
        img = Image.open(io.BytesIO(image_bytes))
        img.thumbnail((max_width, max_height))
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=85)
        buf.seek(0)
        return buf.getvalue()
    except Exception as e:
        print(f"Image resize error: {e}")
        return None


def image_ocr_text(image_bytes):
    try:
        # Convert to grayscale
        nparr = np.frombuffer(image_bytes, np.uint8)
        img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

        # Apply thresholding for better contrast
        _, thresh = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY)

        # Convert back to PIL for pytesseract
        pil_img = Image.fromarray(thresh)

        # Run OCR with language set
        text = pytesseract.image_to_string(pil_img, lang="eng")
        return text.strip()
    except Exception as e:
        print(f"OCR error: {e}")
        return ""


# ----------------------- Document Endpoints -----------------------
@app.route("/doc/pdf_to_word", methods=["POST"])
def pdf_to_word_route():
    file = request.files.get("file")
    if not file:
        return jsonify({"error": "No file provided"}), 400

    pdf_bytes = file.read()
    word_bytes = pdf_to_word_bytes(pdf_bytes)
    if not word_bytes:
        return jsonify({"error": "Conversion failed"}), 500

    key = f"document_tools/{file.filename.rsplit('.', 1)[0]}.docx"
    url = upload_bytes_to_s3(word_bytes, key, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    if not url:
        return jsonify({"error": "Upload failed"}), 500

    return jsonify({"download_url": url})


@app.route("/doc/word_to_pdf", methods=["POST"])
def word_to_pdf_route():
    file = request.files.get("file")
    if not file:
        return jsonify({"error": "No file provided"}), 400

    pdf_bytes = word_to_pdf_bytes(file.read())
    if not pdf_bytes:
        return jsonify({"error": "Conversion failed"}), 500

    key = f"document_tools/{file.filename.rsplit('.', 1)[0]}.pdf"
    url = upload_bytes_to_s3(pdf_bytes, key, "application/pdf")

    if not url:
        return jsonify({"error": "Upload failed"}), 500

    return jsonify({"download_url": url})


@app.route("/doc/image_to_pdf", methods=["POST"])
def image_to_pdf_route():
    file = request.files.get("file")
    if not file:
        return jsonify({"error": "No file provided"}), 400

    pdf_bytes = image_to_pdf_bytes(file.read())
    if not pdf_bytes:
        return jsonify({"error": "Conversion failed"}), 500

    key = f"document_tools/{file.filename.rsplit('.', 1)[0]}.pdf"
    url = upload_bytes_to_s3(pdf_bytes, key, "application/pdf")

    if not url:
        return jsonify({"error": "Upload failed"}), 500

    return jsonify({"download_url": url})


@app.route("/doc/pdf_to_images", methods=["POST"])
def pdf_to_images_route():
    file = request.files.get("file")
    if not file:
        return jsonify({"error": "No file provided"}), 400

    images_bytes = pdf_to_images_bytes(file.read())
    if not images_bytes:
        return jsonify({"error": "Conversion failed"}), 500

    urls = []
    for i, b in enumerate(images_bytes):
        key = f"document_tools/{file.filename.rsplit('.', 1)[0]}_page_{i + 1}.png"
        url = upload_bytes_to_s3(b, key, "image/png")
        if url:
            urls.append(url)

    if not urls:
        return jsonify({"error": "Upload failed"}), 500

    return jsonify({"download_urls": urls})


@app.route("/doc/merge_pdfs", methods=["POST"])
def merge_pdfs_route():
    files = request.files.getlist("files")
    if not files or len(files) < 2:
        return jsonify({"error": "At least two files required"}), 400

    pdf_bytes_list = [f.read() for f in files]
    merged_bytes = merge_pdf_bytes(pdf_bytes_list)

    if not merged_bytes:
        return jsonify({"error": "Merge failed"}), 500

    key = f"document_tools/merged_{int(time.time())}.pdf"
    url = upload_bytes_to_s3(merged_bytes, key, "application/pdf")

    if not url:
        return jsonify({"error": "Upload failed"}), 500

    return jsonify({"download_url": url})


# ----------------------- Image Routes -----------------------
@app.route("/img/resize", methods=["POST"])
def resize_route():
    file = request.files.get("file")
    if not file:
        return jsonify({"error": "No file provided"}), 400

    resized_bytes = resize_image(file.read(), 800, 800)
    if not resized_bytes:
        return jsonify({"error": "Resize failed"}), 500

    key = f"image_tools/{file.filename.rsplit('.', 1)[0]}_resized.jpg"
    url = upload_bytes_to_s3(resized_bytes, key, "image/jpeg")

    if not url:
        return jsonify({"error": "Upload failed"}), 500

    return jsonify({"download_url": url})


@app.route("/img/ocr", methods=["POST"])
def image_ocr_route():
    file = request.files.get("file")
    if not file:
        return jsonify({"error": "No file provided"}), 400

    text = image_ocr_text(file.read())
    return jsonify({"extracted_text": text})


# ----------------------- Combined / AI Endpoints -----------------------
@app.route("/YcloudSpark", methods=["POST"])
def YcloudSpark():
    files = request.files.getlist("files")
    question = request.form.get("question")
    gq = request.form.get("quiz")
    flashcards = request.form.get("flashcards")

    if not files:
        return jsonify({"error": "Please select a file"}), 400

    # Extract all text
    all_text = ""
    for f in files:
        try:
            all_text += extract(f) + "\n"
        except Exception as e:
            print("Error extracting file:", e)

    result = {}

    # Summarization
    summary_prompt = f"Extract brief key points perfectly from this text so that everyone can understand:\n{all_text}"
    result["summary"] = ask_ai(summary_prompt)

    # Q&A
    if question:
        qa_prompt = f"Based on this text, answer the question in a simple way:\n{all_text}\nQuestion: {question}"
        result["answer"] = ask_ai(qa_prompt, max_tokens=200)

    # Quiz
    if gq and gq.lower() == "yes":
        quiz_prompt = f"Generate 5 multiple-choice quiz questions (2 difficult, 2 easy, 1 moderate) with answers from this text:\n{all_text}"
        result["quiz"] = ask_ai(quiz_prompt)

    # Flashcards
    if flashcards and flashcards.lower() == "yes":
        flash_prompt = f"Create flashcards (Q&A) from this text. Include both easy and difficult level questions:\n{all_text}"
        result["flashcards"] = ask_ai(flash_prompt)

    return jsonify(result)


@app.route("/ai/resume_builder", methods=["POST"])
def resume_builder():
    text_input = request.form.get("text")
    file = request.files.get("file")

    if not text_input and not file:
        return jsonify({"error": "Please provide text or upload a file"}), 400

    if file and not text_input:
        try:
            if file.filename.lower().endswith(".pdf"):
                text_input = extract(file)
            elif file.filename.lower().endswith((".docx", ".doc")):
                doc = Document(file)
                text_input = "\n".join([p.text for p in doc.paragraphs])
            else:
                text_input = file.read().decode("utf-8")
        except Exception as e:
            return jsonify({"error": f"File processing failed: {e}"}), 500

    prompt = f"Create a professional, ATS-friendly resume based on the following information and you know that resume should be a very clear format:\n{text_input}"
    resume_text = ask_ai(prompt, max_tokens=800)
    return jsonify({"resume": resume_text})


@app.route("/ai/mindmap", methods=["POST"])
def mindmap_generator():
    text_input = request.form.get("text")
    file = request.files.get("file")

    if not text_input and not file:
        return jsonify({"error": "Please provide text or upload a file"}), 400

    if file and not text_input:
        try:
            if file.filename.lower().endswith(".pdf"):
                text_input = extract(file)
            elif file.filename.lower().endswith((".docx", ".doc")):
                doc = Document(file)
                text_input = "\n".join([p.text for p in doc.paragraphs])
            else:
                text_input = file.read().decode("utf-8")
        except Exception as e:
            return jsonify({"error": f"File processing failed: {e}"}), 500

    prompt = (
        f"Generate a detailed hierarchical mindmap in JSON format based on the following text:\n{text_input}\n"
        f"Explain topics in a way that a small child can also understand, with nodes and children."
    )
    mindmap_json = ask_ai(prompt, max_tokens=800)
    return jsonify({"mindmap": mindmap_json})


# ----------------------- Voice-to-Document -----------------------
@app.route("/ai/voice_to_document", methods=["POST"])
def voice_to_document():
    file = request.files.get("file")
    if not file:
        return jsonify({"error": "Please upload an audio file"}), 400

    try:
        temp_path = os.path.join(UPLOAD_FOLDER, file.filename)
        file.save(temp_path)

        with open(temp_path, "rb") as audio_file:
            transcription = client.audio.transcriptions.create(
                model="whisper-1",
                file=audio_file
            )

        text_output = transcription.text
        key = f"voice_to_doc/{file.filename.rsplit('.', 1)[0]}.txt"
        url = upload_bytes_to_s3(text_output.encode("utf-8"), key, "text/plain")

        os.remove(temp_path)

        return jsonify({"transcription": text_output, "download_url": url})
    except Exception as e:
        print("Voice-to-Document Error:", e)
        return jsonify({"error": f"Transcription failed: {e}"}), 500


if __name__ == "__main__":
    # change debug=False in production
    app.run(host="0.0.0.0", port=4000, debug=True)
