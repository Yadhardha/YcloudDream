from flask import Flask, request, jsonify, send_file, render_template, redirect,send_from_directory
from flask_cors import CORS
from werkzeug.utils import secure_filename
import os
import random
import string
import io
import qrcode
import time
import threading
from dotenv import load_dotenv
from PyPDF2 import PdfReader, PdfMerger
from docx import Document
from reportlab.pdfgen import canvas
from PIL import Image
from pdf2image import convert_from_bytes
import pytesseract
import google.generativeai as genai
from google.generativeai import models
import cv2
import numpy as np
from google.generativeai import GenerativeModel
import re
from fpdf import FPDF # type: ignore
import uuid
from PIL import Image
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import letter
import pdfplumber
from docx import Document
import pytesseract
from pdf2image import convert_from_bytes
import fitz


# ----------------------- App Setup -----------------------
app = Flask(__name__)
CORS(app)
app.config['MAX_CONTENT_LENGTH'] = 512 * 1024 * 1024  # 512 MB
app.config["PROPAGATE_EXCEPTIONS"] = True
app.config["DEBUG"] = True


UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "outputs"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)
app.config['OUTPUT_FOLDER'] = OUTPUT_FOLDER
app.config['UPLOAD_FOLDER']= UPLOAD_FOLDER
DELETE_AFTER = 2 * 60 * 60
CHECK_INTERVAL = 300


# load env file


load_dotenv()

try:
    gemini_api_key = os.environ.get("GOOGLE_API_KEY")
    if not gemini_api_key:
        raise ValueError("GOOGLE_API_KEY not found in .env file")
    genai.configure(api_key=gemini_api_key)
    print("✅ Google Generative AI configured successfully!")
except Exception as e:
    print(f"❌ Error configuring Google Generative AI: {e}")


def generate_code(length=6):
    return ''.join(random.choices(string.ascii_uppercase + string.digits, k=length))
def auto_delete():
    while True:
        now = time.time()
        if os.path.exists(UPLOAD_FOLDER):
            for i in os.listdir(UPLOAD_FOLDER):
                path = os.path.join(UPLOAD_FOLDER, i)
                if os.path.isfile(path) and now - os.path.getmtime(path) > DELETE_AFTER:
                    try:
                        os.remove(path)
                    except Exception as e:
                        print("Error deleting file:", e)
        time.sleep(CHECK_INTERVAL)


threading.Thread(target=auto_delete, daemon=True).start()

# ----------------------- Routes -----------------------
@app.route("/")
def home():
    return render_template("index.html")

# ------------------------ Upload and receiver section ------------------------#

@app.route("/upload", methods=["POST"])
def upload():
    file = request.files.get("file")
    if not file or file.filename == "":
        return jsonify({"error": "No file uploaded"}), 400
    
    try:
        filename, filepath = save_file(file)
        download_url = f"{request.url_root}download/{filename}"
        return jsonify({"code": filename.split("_")[0], "filename": filename, "download_url": download_url})
    except Exception as e:
        return jsonify({"error": f"Upload failed: {e}"}), 500


@app.route("/api/status", methods=["GET"])
def api_status():
    return jsonify({
        "status": "running",
        "timestamp": time.time()
    }), 200


# THIS IS THE CORRECT, UNIVERSAL DOWNLOAD ROUTE
@app.route("/download/<path:filepath>")
def download(filepath):
    # Determine the base folder based on the path
    if "document_tools" in filepath or "image_tools" in filepath:
        base_folder = OUTPUT_FOLDER
    else:
        base_folder = UPLOAD_FOLDER
    
    # Securely build the full path
    full_path = os.path.join(base_folder, filepath)
    directory = os.path.dirname(full_path)
    filename = os.path.basename(full_path)

    if not os.path.exists(full_path):
         return jsonify({"error": "File not found or has expired"}), 404

    return send_from_directory(directory, filename, as_attachment=True)


def save_file(file):
    code = generate_code()
    original_filename = secure_filename(file.filename)
    new_filename = f"{code}_{original_filename}"
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], new_filename)
    file.save(file_path)
    return new_filename, file_path


@app.route("/check_file/<code>", methods=["GET"])
def check_file(code):
    for fname in os.listdir(app.config['UPLOAD_FOLDER']):
        if fname.startswith(code + "_"):
            return jsonify({"filename": fname})
    return jsonify({"error": "File not found"}), 404


@app.route("/get_qr/<code>", methods=["GET"])
def get_qr(code):
    filename = None
    for f in os.listdir(app.config['UPLOAD_FOLDER']):
        if f.startswith(code + "_"):
            filename = f
            break
    if not filename:
        return jsonify({"error": "File not found"}), 404
    
    PUBLIC_URL = os.getenv("PUBLIC_BACKEND_URL")
    download_url = f"{PUBLIC_URL}/download/{filename}"
    
    qr = qrcode.QRCode(version=1, box_size=10, border=5)
    qr.add_data(download_url)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    
    return send_file(buf, mimetype='image/png')


def save_file_locally(file_bytes, filename, subfolder=""):
    """Saves file bytes to the OUTPUT_FOLDER with an optional subfolder."""
    try:
        # Create the full directory path if it doesn't exist
        directory = os.path.join(OUTPUT_FOLDER, subfolder)
        os.makedirs(directory, exist_ok=True)
        
        filepath = os.path.join(directory, secure_filename(filename))
        with open(filepath, "wb") as f:
            f.write(file_bytes)
            
        # Return the relative path for the download URL
        return os.path.join(subfolder, filename)
    except Exception as e:
        print(f"Error saving file locally: {e}")
        return None
def save_text_as_pdf(text, filename):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)
    for line in text.split("\n"):
        pdf.multi_cell(0, 10, line)
    pdf.output(os.path.join("outputs", filename))

def save_text_as_docx(text, filename):
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(os.path.join("outputs", filename))

# ----------------------- Document Helpers -----------------------
def extract(file_like):

    """
    Accepts a file-like object (supports .read()). Returns extracted text (PDF).
    """
    try:
        # if it's a Flask FileStorage, read bytes
        if hasattr(file_like, "read"):
            data = file_like.read()
        else:
            data = file_like
        reader = PdfReader(io.BytesIO(data))
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text
    except Exception as e:
        print(f"PDF extraction error: {e}")
        return ""


def pdf_to_word_bytes(pdf_bytes):
    try:
        doc = Document()

        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()

                # OCR fallback
                if not text or text.strip() == "":
                    images = convert_from_bytes(pdf_bytes, first_page=i+1, last_page=i+1)
                    for img in images:
                        text = pytesseract.image_to_string(img)

                doc.add_paragraph(f"--- Page {i + 1} ---")
                doc.add_paragraph(text if text else "")

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    except Exception as e:
        print("PDF → Word error:", e)
        return None



def word_to_pdf_bytes(word_bytes):
    try:
        docx = Document(io.BytesIO(word_bytes))

        buffer = io.BytesIO()
        pdf = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        for para in docx.paragraphs:
            text = para.text.strip()
            if text:
                story.append(Paragraph(text, styles["Normal"]))

        pdf.build(story)
        buffer.seek(0)
        return buffer.getvalue()

    except Exception as e:
        print("Word → PDF error:", e)
        return None

def image_to_pdf_bytes(image_bytes):
    try:
        img = Image.open(io.BytesIO(image_bytes))

        if img.mode != "RGB":
            img = img.convert("RGB")

        buffer = io.BytesIO()
        img.save(buffer, "PDF", resolution=300.0)
        buffer.seek(0)
        return buffer.getvalue()

    except Exception as e:
        print("Image → PDF error:", e)
        return None
# PyMuPDF

def pdf_to_images_bytes(pdf_bytes):
    try:
        pdf = fitz.open(stream=pdf_bytes, filetype="pdf")
        images = []

        for page_no in range(len(pdf)):
            page = pdf.load_page(page_no)
            pix = page.get_pixmap(dpi=300)
            images.append(pix.tobytes("png"))

        return images

    except Exception as e:
        print(f"PDF to images conversion error: {e}")
        return []


def merge_pdf_bytes(pdf_files_bytes):
    try:
        merger = PdfMerger()
        for pdf_bytes in pdf_files_bytes:
            merger.append(io.BytesIO(pdf_bytes))
        out = io.BytesIO()
        merger.write(out)
        merger.close()
        out.seek(0)
        return out.getvalue()
    except Exception as e:
        print(f"PDF merge error: {e}")
        return None


def resize_image(image_bytes, max_width=800, max_height=800):
    try:
        img = Image.open(io.BytesIO(image_bytes))
        img.thumbnail((max_width, max_height))
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=85)
        buf.seek(0)
        return buf.getvalue()
    except Exception as e:
        print(f"Image resize error: {e}")
        return None


def image_ocr_text(image_bytes):
    try:
        # Convert to grayscale
        nparr = np.frombuffer(image_bytes, np.uint8)
        img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

        # Apply thresholding for better contrast
        _, thresh = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY)

        # Convert back to PIL for pytesseract
        pil_img = Image.fromarray(thresh)

        # Run OCR with language set
        text = pytesseract.image_to_string(pil_img, lang="eng")
        return text.strip()
    except Exception as e:
        print(f"OCR error: {e}")
        return ""


# ----------------------- Document Endpoints -----------------------
@app.route("/doc/pdf_to_word", methods=["POST"])
def pdf_to_word_route():
    file = request.files.get("file")
    if not file:
        return jsonify({"error": "No file provided"}), 400

    pdf_bytes = file.read()
    word_bytes = pdf_to_word_bytes(pdf_bytes)
    if not word_bytes:
        return jsonify({"error": "Conversion failed"}), 500

    filename = f"{file.filename.rsplit('.', 1)[0]}.docx"
    save_file_locally(word_bytes, filename, subfolder="document_tools")
    url = f"/download/document_tools/{filename}"


    if not url:
        return jsonify({"error": "Upload failed"}), 500

    return jsonify({"download_url": url})


@app.route("/doc/word_to_pdf", methods=["POST"])
def word_to_pdf_route():
    file = request.files.get("file")
    if not file:
        return jsonify({"error": "No file provided"}), 400

    pdf_bytes = word_to_pdf_bytes(file.read())
    if not pdf_bytes:
        return jsonify({"error": "Conversion failed"}), 500

    filename = f"{file.filename.rsplit('.', 1)[0]}.pdf"
    save_file_locally(pdf_bytes, filename, subfolder="document_tools")
    url = f"/download/document_tools/{filename}"


    if not url:
        return jsonify({"error": "Upload failed"}), 500

    return jsonify({"download_url": url})


@app.route("/doc/image_to_pdf", methods=["POST"])
def image_to_pdf_route():
    file = request.files.get("file")
    if not file:
        return jsonify({"error": "No file provided"}), 400

    pdf_bytes = image_to_pdf_bytes(file.read())
    if not pdf_bytes:
        return jsonify({"error": "Conversion failed"}), 500

    # Save locally
    filename = f"{file.filename.rsplit('.', 1)[0]}.pdf"
    save_file_locally(pdf_bytes, filename, subfolder="document_tools")
    url = f"/download/document_tools/{filename}"

    return jsonify({"download_url": url})


@app.route("/doc/pdf_to_images", methods=["POST"])
def pdf_to_images_route():
    file = request.files.get("file")
    if not file:
        return jsonify({"error": "No file provided"}), 400

    images_bytes = pdf_to_images_bytes(file.read())
    if not images_bytes:
        return jsonify({"error": "Conversion failed"}), 500

    urls = []
    for i, b in enumerate(images_bytes):
        filename = f"{file.filename.rsplit('.', 1)[0]}_page_{i+1}.png"
        save_file_locally(b, filename, subfolder="document_tools")
        urls.append(f"/download/document_tools/{filename}")


    if not urls:
        return jsonify({"error": "Upload failed"}), 500

    return jsonify({"download_urls": urls})


@app.route("/doc/merge_pdfs", methods=["POST"])
def merge_pdfs_route():
    files = request.files.getlist("files")
    if not files or len(files) < 2:
        return jsonify({"error": "At least two files required"}), 400

    pdf_bytes_list = [f.read() for f in files]
    merged_bytes = merge_pdf_bytes(pdf_bytes_list)

    if not merged_bytes:
        return jsonify({"error": "Merge failed"}), 500

    filename = f"merged_{int(time.time())}.pdf"
    save_file_locally(merged_bytes, filename, subfolder="document_tools")
    url = f"/download/document_tools/{filename}"


    if not url:
        return jsonify({"error": "Upload failed"}), 500

    return jsonify({"download_url": url})


# ----------------------- Image Routes -----------------------
@app.route("/img/resize", methods=["POST"])
def resize_route():
    file = request.files.get("file")
    if not file:
        return jsonify({"error": "No file provided"}), 400

    resized_bytes = resize_image(file.read(), 800, 800)
    if not resized_bytes:
        return jsonify({"error": "Resize failed"}), 500

    filename = f"{file.filename.rsplit('.', 1)[0]}_resized.jpg"
    save_file_locally(resized_bytes, filename, subfolder="image_tools")
    url = f"/download/image_tools/{filename}"


    if not url:
        return jsonify({"error": "Upload failed"}), 500

    return jsonify({"download_url": url})


@app.route("/img/ocr", methods=["POST"])
def image_ocr_route():
    file = request.files.get("file")
    if not file:
        return jsonify({"error": "No file provided"}), 400

    text = image_ocr_text(file.read())
    return jsonify({"extracted_text": text})

def ask_Ycloudai(prompt, max_tokens=2048):
    try:
        model = genai.GenerativeModel("gemini-2.0-flash")

        response = model.generate_content(
            prompt,
            generation_config={"max_output_tokens": max_tokens}
        )

        # Safety check
        if not response.candidates or not response.candidates[0].content.parts:
            return "⚠️ AI blocked the response due to safety filters."

        return response.text.strip()

    except Exception as e:
        print("Gemini API error:", e)
        return f"⚠️ Error: {str(e)}"

def _mime_from_filename(filename):
    fn = filename.lower()

    if fn.endswith(".pdf"):
        return "application/pdf"

    if fn.endswith(".docx"):
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    if fn.endswith(".txt"):
        return "text/plain"

    if fn.endswith(".png"):
        return "image/png"

    if fn.endswith(".jpg") or fn.endswith(".jpeg"):
        return "image/jpeg"

    if fn.endswith(".webm") or fn.endswith(".opus"):
        return "audio/webm"



    if fn.endswith(".wav"):
        return "audio/wav"

    if fn.endswith(".mp3"):
        return "audio/mpeg"

    # fallback for unknown/extensions
    return "application/octet-stream"


def extract_text_from_file(file):
    try:
        filename = file.filename.lower()
        file_bytes = file.read()

        print("\n🟦 [TRACE] Extracting file:", filename)

        # -------------------------
        # 📄 DOCX extraction
        # -------------------------
        if filename.endswith(".docx"):
            print("🟩 [OK] DOCX detected — extracting using python-docx")
            doc = Document(io.BytesIO(file_bytes))
            text = "\n".join([p.text for p in doc.paragraphs])
            return text, None

        # -------------------------
        # 📄 PDF extraction
        # -------------------------
        if filename.endswith(".pdf"):
            print("🟩 [OK] PDF detected — extracting using pdfplumber")
            text = ""
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text, None

        # -------------------------
        # 📄 TXT extraction
        # -------------------------
        if filename.endswith(".txt"):
            print("🟩 [OK] TXT detected")
            return file_bytes.decode("utf-8"), None

        print("🟥 [ERROR] Unsupported file:", filename)
        return None, "Unsupported file format."

    except Exception as e:
        print("🟥 [FATAL] Extraction Error:")
        traceback.print_exc()
        return None, f"Extract failed: {str(e)}"


def chunk_text(text, chunk_size=4000, overlap=200):
    """Splits a long text into smaller, overlapping chunks."""
    if not text:
        return []
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start += chunk_size - overlap
    return chunks

# ----------------------- Ycloud Spark - Document Processor -----------------------
def generate_summary(file_bytes, filename):
    mime = _mime_from_filename(filename)

    prompt = (
        "Read the attached document fully and generate a detailed structured summary. "
        "Include an executive summary, key points, sections, bullet points, and important insights."
    )

    model = genai.GenerativeModel("gemini-2.0-flash")

    response = model.generate_content([
        prompt,
        {"mime_type": mime, "data": file_bytes}
    ])

    if not response.candidates or not response.candidates[0].content.parts:
        return "⚠️ AI blocked the summary due to safety."

    return response.text



def generate_qa(file_bytes, filename, question):
    mime = _mime_from_filename(filename)

    prompt = (
        f"Read the attached document and answer this question based ONLY on the document:\n\n"
        f"Question: {question}"
    )

    model = genai.GenerativeModel("gemini-2.0-flash")

    response = model.generate_content([
        prompt,
        {"mime_type": mime, "data": file_bytes}
    ])

    if not response.candidates or not response.candidates[0].content.parts:
        return "⚠️ AI blocked the answer due to safety."

    return response.text

  
def generate_quiz(file_bytes, filename):
    mime = _mime_from_filename(filename)

    prompt = (
        "Read the attached document and generate 5 multiple-choice questions "
        "(2 easy, 2 medium, 1 hard). Include answers."
    )

    model = genai.GenerativeModel("gemini-2.0-flash")

    response = model.generate_content([
        prompt,
        {"mime_type": mime, "data": file_bytes}
    ])

    if not response.candidates or not response.candidates[0].content.parts:
        return "⚠️ AI blocked the quiz due to safety."

    return response.text


def generate_flashcards(file_bytes, filename):
    mime = _mime_from_filename(filename)

    prompt = (
        "Read the attached document and generate 5 flashcards (Q & A format)."
    )

    model = genai.GenerativeModel("gemini-2.0-flash")

    response = model.generate_content([
        prompt,
        {"mime_type": mime, "data": file_bytes}
    ])

    if not response.candidates or not response.candidates[0].content.parts:
        return "⚠️ AI blocked the flashcards due to safety."

    return response.text

@app.route("/YcloudSpark", methods=["POST"])
def ycloud_spark_processor():
    try:
        # 1️⃣ Check if files are uploaded
        files = request.files.getlist("files")
        if not files or all(f.filename == '' for f in files):
            return jsonify({"error": "Please select at least one valid file"}), 400

        # 2️⃣ Extract text from files
        all_text = ""
        for file in files:
            try:
                text, error = extract_text_from_file(file)
                if error:
                    return jsonify({"error": f"Error extracting text from {file.filename}: {error}"}), 400
                all_text += text.strip() + "\n"
            except Exception as e:
                return jsonify({"error": f"Unexpected error while reading {file.filename}: {str(e)}"}), 500

        # 3️⃣ Convert all extracted text into bytes
        full_text = all_text.strip()
        if not full_text:
            return jsonify({"error": "No text extracted from files"}), 400

        full_bytes = full_text.encode("utf-8")
        filename = "combined.txt"   # so mime = text/plain

        # 4️⃣ Read optional inputs
        question = request.form.get("question")
        gq = request.form.get("quiz")
        flashcards = request.form.get("flashcards")

        # 5️⃣ Call AI functions correctly
        result = {}

        # ------ Summary ------
        try:
            result["summary"] = generate_summary(full_bytes, filename)
        except Exception as e:
            result["summary_error"] = f"Error generating summary: {str(e)}"

        # ------ Q & A ------
        if question:
            try:
                result["answer"] = generate_qa(full_bytes, filename, question)
            except Exception as e:
                result["answer_error"] = f"Error generating answer: {str(e)}"

        # ------ Quiz ------
        if gq and gq.lower() == "yes":
            try:
                result["quiz"] = generate_quiz(full_bytes, filename)
            except Exception as e:
                result["quiz_error"] = f"Error generating quiz: {str(e)}"

        # ------ Flashcards ------
        if flashcards and flashcards.lower() == "yes":
            try:
                result["flashcards"] = generate_flashcards(full_bytes, filename)
            except Exception as e:
                result["flashcards_error"] = f"Error generating flashcards: {str(e)}"

        return jsonify(result)

    except Exception as e:
        return jsonify({"error": f"Unexpected server error: {str(e)}"}), 500


# ----------------------- Ycloud Mindmap -----------------------
@app.route("/ai/mindmap", methods=["POST"])
def mindmap_generator():
    text_input = request.form.get("text")
    file = request.files.get("file")

    if not text_input and not file:
        return jsonify({"error": "Please provide text or upload a file"}), 400

    if file:
        text_input, error = extract_text_from_file(file)
        if error:
            return jsonify({"error": error}), 400

    # STRICT instruction: Do NOT use emojis. Do NOT use FontAwesome.
    prompt = f"""
Convert the following text into a clean, readable hierarchical **Mermaid mindmap**.

### RULES (IMPORTANT)
- ONLY use valid Mermaid mindmap syntax.
- DO NOT use emojis.
- DO NOT use Font Awesome icons.
- The structure must start with:
mindmap
  root(Your Main Topic)
    Child
      Sub child
    Another child
- Keep the text short, clean, and readable.
- Identify the main topic automatically from the text.

### Text:
{text_input}
"""

    mindmap_syntax = ask_Ycloudai(prompt)

    # Clean result: remove markdown formatting if any
    mindmap_clean = mindmap_syntax.replace("```mermaid", "").replace("```", "").strip()

    return jsonify({"mindmap_code": mindmap_clean})


ROLE_SKILLS = {
    "Data Scientist": [
        "Python", "R", "SQL", "Pandas", "NumPy", "Scikit-learn", "TensorFlow",
        "PyTorch", "Statistics", "Machine Learning", "Deep Learning", "NLP",
        "Data Visualization", "Matplotlib", "Seaborn"
    ],
    "Java Developer": [
        "Java", "Spring Boot", "Hibernate", "REST API", "Microservices", "Maven",
        "JUnit", "Docker", "SQL", "Git", "Eclipse", "IntelliJ", "OOP"
    ],
    "Python Developer": [
        "Python", "Django", "Flask", "FastAPI", "SQLAlchemy", "REST API",
        "Pandas", "NumPy", "Pytest", "Git", "Docker", "OOP"
    ],
    "Cloud DevOps Engineer": [
        "AWS", "Azure", "GCP", "Terraform", "Ansible", "Jenkins", "Docker",
        "Kubernetes", "Linux", "Bash", "Python", "Git", "CI/CD", "Monitoring",
        "CloudFormation"
    ],
    "DevOps Engineer": [
        "Jenkins", "Docker", "Kubernetes", "Ansible", "Terraform", "Linux",
        "Bash", "Git", "AWS", "CI/CD", "Monitoring", "Networking", "Python"
    ],
    "Data Analyst": [
        "Excel", "SQL", "Python", "Power BI", "Tableau", "Pandas", "NumPy",
        "Statistics", "Data Visualization", "Data Cleaning"
    ],
    "Data Engineer": [
        "Python", "SQL", "Spark", "Hadoop", "Kafka", "AWS", "GCP", "Azure",
        "ETL", "Airflow", "Databricks", "BigQuery", "Redshift", "Data Warehousing"
    ],
    "SAP Developer": [
        "SAP ABAP", "SAP HANA", "SAP Fiori", "SAP UI5", "SAP BW", "SAP NetWeaver",
        "OData Services", "SAP BusinessObjects"
    ],
    "Prompt Engineer": [
        "LLMs", "OpenAI", "HuggingFace", "LangChain", "Python", "NLP",
        "Prompt Design", "Fine-tuning", "RAG", "Vector Databases", "JSON"
    ]
}


def ats_score(resume_text: str, role: str):
    text = resume_text.lower()
    skills = ROLE_SKILLS.get(role, [])
    total = len(skills)
    threshold = total / 2

    matched = []
    for skill in skills:
        skill_pattern = re.escape(skill.lower())
        if re.search(skill_pattern, text):
            matched.append(skill)

    missing = [s for s in skills if s not in matched]
    M = len(matched)

    # Threshold scoring
    if M < threshold:
        score = (M / total) * 50
    else:
        score = 50 + ((M - threshold) / threshold) * 50

    return round(score, 2), matched, missing


def gemini_feedback(resume_text: str, role: str, matched: list, missing: list, score: float):
    prompt = f"""
You are an expert ATS Resume Analyzer.

Candidate Role: {role}

Resume Score: {score}%
Matched Skills: {matched}
Missing Skills: {missing}

Resume Text (first 1500 chars):
{resume_text[:1500]}

Provide output in this structure:

Highlights:
- ...

Suggestions:
- ...
"""

    model = genai.GenerativeModel("gemini-2.0-flash")
    response = model.generate_content(prompt)

    return response.text


@app.route("/ai/resume_builder", methods=["POST"])
def resume_builder():
    try:
        print("\n==============================")
        print("🟦 [TRACE] Resume Analyzer Triggered")
        print("==============================")

        file = request.files.get("file")
        selected_role = request.form.get("role")

        print("🟦 [TRACE] Role selected:", selected_role)
        print("🟦 [TRACE] File received:", file.filename if file else None)

        if not file:
            print("🟥 [ERROR] No file uploaded")
            return jsonify({"error": "❌ No file uploaded"}), 400

        if not selected_role:
            print("🟥 [ERROR] No role selected")
            return jsonify({"error": "❌ No role selected"}), 400

        print("🟦 [TRACE] Extracting resume text...")
        text_input, error = extract_text_from_file(file)

        if error:
            print("🟥 [ERROR] Extract error:", error)
            return jsonify({"error": error}), 400

        if not text_input.strip():
            print("🟥 [ERROR] Extracted text was empty")
            return jsonify({"error": "❌ Could not read text from resume"}), 400

        print("🟦 [TRACE] Running ATS score...")
        score, matched, missing = ats_score(text_input, selected_role)

        print("🟩 [TRACE] ATS Score Done:", score)
        print("🟦 [TRACE] Matched:", matched)
        print("🟦 [TRACE] Missing:", missing)

        print("🟦 [TRACE] Asking Gemini for feedback...")
        ai_feedback = gemini_feedback(text_input, selected_role, matched, missing, score)

        print("🟩 [TRACE] Gemini feedback received")

        return jsonify({
            "resume_score": score,
            "matched_skills": matched,
            "missing_skills": missing,
            "ai_feedback": ai_feedback
        })

    except Exception as e:
        print("🟥 [FATAL] EXCEPTION in resume_builder")
        traceback.print_exc()  # 🔥 FULL TRACEBACK HERE
        return jsonify({"error": f"🔥 Internal Error: {str(e)}"}), 500


def ask_Ycloudai_audio(file_bytes, mime_type):
    try:
        model = genai.GenerativeModel("gemini-2.0-flash")

        prompt = (
            "Transcribe this audio clearly. "
            "Fix mistakes and punctuation. "
            "Do NOT summarize. "
            "Return clean readable text."
        )

        response = model.generate_content([
            prompt,
            {"mime_type": mime_type, "data": file_bytes}
        ])

        return response.text.strip()

    except Exception as e:
        print("Audio transcription error:", e)
        return None




@app.route("/ai/voice_to_document", methods=["POST"])
def voice_to_document():
    try:
        file = request.files.get("file")
        if not file:
            return jsonify({"error": "Please upload an audio file"}), 400
        
        filename = file.filename
        file_bytes = file.read()

        # Detect mime
        mime = _mime_from_filename(filename)
        if mime not in ["audio/webm", "audio/wav", "audio/mpeg"]:
            return jsonify({"error": "Unsupported audio format"}), 400

        # Transcribe using Gemini
        text_output = ask_Ycloudai_audio(file_bytes, mime)
        if not text_output:
            return jsonify({"error": "AI transcription failed"}), 500

        output_format = request.form.get("format", "pdf").lower()

        # Save as PDF
        if output_format == "pdf":
            output_file = f"{uuid.uuid4()}.pdf"
            pdf_doc = FPDF()
            pdf_doc.add_page()
            pdf_doc.set_font("Arial", size=12)
            pdf_doc.multi_cell(0, 10, text_output)
            pdf_bytes = pdf_doc.output(dest="S").encode("latin1")
            saved_path = save_file_locally(pdf_bytes, output_file, subfolder="document_tools")

        # Save as DOCX
        else:
            output_file = f"{uuid.uuid4()}.docx"
            doc = Document()
            doc.add_paragraph(text_output)
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            saved_path = save_file_locally(buffer.getvalue(), output_file, subfolder="document_tools")

        return jsonify({
            "text": text_output,
            "download_url": f"/download/{saved_path}"
        })

    except Exception as e:
        print("Voice-to-Document Error:", e)
        return jsonify({"error": str(e)}), 500





if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)
